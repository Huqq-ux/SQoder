import os
import asyncio
import logging
import re
from fastapi import APIRouter, UploadFile, File, Query, HTTPException, BackgroundTasks
from Coder.server.schemas import KnowledgeSearchRequest

logger = logging.getLogger(__name__)
router = APIRouter()

_SAFE_FILENAME_RE = re.compile(r'^[\w\-\.]+$')
_ALLOWED_SUFFIXES = {".txt", ".md", ".pdf", ".docx", ".pptx", ".xlsx", ".csv", ".epub"}

_INDEX_BASE = os.path.join(os.path.dirname(__file__), "..", "..", "knowledge", "index")


async def _index_and_extract(course_id: str, filepath: str, safe_name: str, ext: str,
                              content_len: int, chunks: list, doc: dict):
    """Background task: add to vector store and extract knowledge points."""
    from Coder.knowledge.vector_store import VectorStore
    from Coder.storage.course_manager import CourseManager
    from Coder.knowledge.text_splitter import StructuredTextSplitter

    try:
        if course_id:
            store = VectorStore(store_path=os.path.join(_INDEX_BASE, course_id))
            store.add_documents(chunks)
            await CourseManager.register_file(
                course_id, safe_name, ext, content_len, len(chunks),
                os.path.join(_INDEX_BASE, course_id),
            )
        else:
            store = VectorStore()
            store.add_documents(chunks)
            await CourseManager.register_file(
                None, safe_name, ext, content_len, len(chunks),
                os.path.join(_INDEX_BASE, ""),
            )
        logger.info(f"向量索引完成: {safe_name}")

        # Extract knowledge points via LLM
        if course_id:
            try:
                await CourseManager.delete_knowledge_points_by_file(course_id, safe_name)
                raw_content = doc.get("content", "")
                raw_metadata = doc.get("metadata", {})
                splitter = StructuredTextSplitter()
                kps = splitter.extract_knowledge_points(raw_content, raw_metadata)
                if kps:
                    n = await CourseManager.batch_add_knowledge_points(course_id, kps)
                    logger.info(f"已提取 {n} 个知识点: {safe_name}")
            except Exception as e:
                logger.warning(f"知识点提取失败: {safe_name}: {e}")
    except Exception as e:
        logger.error(f"后台索引失败: {safe_name}: {e}")


def _get_course_store(course_id: str):
    """Get a VectorStore for a specific course."""
    from Coder.knowledge.vector_store import VectorStore
    store_path = os.path.join(_INDEX_BASE, course_id)
    return VectorStore(store_path=store_path)


def _get_global_store():
    """Get the global (legacy) VectorStore."""
    from Coder.knowledge.vector_store import VectorStore
    return VectorStore()


async def _resolve_course(identifier: str) -> dict | None:
    """Resolve a course identifier (slug or UUID) to a course dict."""
    from Coder.storage.course_manager import CourseManager
    course = await CourseManager.get_course(identifier)
    if not course:
        course = await CourseManager.get_course_by_slug(identifier)
    return course


@router.post("/upload")
async def upload_documents(
    files: list[UploadFile] = File(...),
    course_id: str = Query(default=""),
    background_tasks: BackgroundTasks = None,
):
    """Upload documents. File saved immediately; indexing runs in background."""
    docs_dir = os.path.join(
        os.path.dirname(__file__), "..", "..", "knowledge", "docs"
    )
    docs_dir = os.path.normpath(docs_dir)
    os.makedirs(docs_dir, exist_ok=True)

    from Coder.knowledge.document_loader import DocumentLoader
    from Coder.knowledge.text_splitter import StructuredTextSplitter
    from Coder.storage.course_manager import CourseManager

    loader = DocumentLoader()
    splitter = StructuredTextSplitter()

    results = []
    for file in files:
        safe_name = os.path.basename(file.filename or "unknown")
        ext = os.path.splitext(safe_name)[1].lower()
        if ext not in _ALLOWED_SUFFIXES:
            results.append({"filename": safe_name, "chunks": 0, "status": f"unsupported format: {ext}"})
            continue

        content = await file.read()
        filepath = os.path.join(docs_dir, safe_name)
        with open(filepath, "wb") as f:
            f.write(content)

        try:
            doc = loader.load(filepath)
            chunks = splitter.split_documents([doc])

            if course_id:
                course = await _resolve_course(course_id)
                if not course:
                    results.append({"filename": safe_name, "chunks": 0, "status": "课程不存在"})
                    continue
                cid = course["id"]
            else:
                cid = ""

            # Schedule background indexing + knowledge point extraction
            background_tasks.add_task(
                _index_and_extract, cid, filepath, safe_name, ext,
                len(content), chunks, doc,
            )

            results.append({"filename": safe_name, "chunks": len(chunks), "status": "indexing"})
        except Exception as e:
            logger.error(f"Upload failed for {safe_name}: {e}")
            results.append({"filename": safe_name, "chunks": 0, "status": f"error: {e}"})

    return {"results": results}


@router.post("/search")
async def search_knowledge(req: KnowledgeSearchRequest, course_id: str = Query(default="")):
    from Coder.knowledge.retriever import Retriever
    from Coder.knowledge.vector_store import VectorStore

    if course_id:
        course = await _resolve_course(course_id)
        if not course:
            return {"results": [], "available": False}
        store = _get_course_store(course["id"])
    else:
        store = _get_global_store()

    if not store.is_available():
        return {"results": [], "available": False}

    docs = store.similarity_search(req.query, k=req.k)
    results = []
    for doc in docs:
        results.append({
            "content": doc.page_content[:500],
            "metadata": {
                "filename": doc.metadata.get("filename", ""),
                "section": doc.metadata.get("section", ""),
                "relevance_score": doc.metadata.get("relevance_score", 0),
            },
        })
    return {"results": results, "available": True}


@router.get("/documents")
async def list_knowledge_documents(course_id: str = Query(default="")):
    """List uploaded documents, optionally filtered by course."""
    if course_id:
        course = await _resolve_course(course_id)
        if not course:
            return {"documents": []}
        from Coder.storage.course_manager import CourseManager
        docs = await CourseManager.list_files(course["id"])
        return {"documents": [{
            "id": d["id"],
            "filename": d["filename"],
            "file_type": d["file_type"],
            "size": d["file_size"],
            "chunks": d["chunk_count"],
            "course_slug": course.get("slug", ""),
            "course_name": course.get("name", ""),
            "status": "indexed",
        } for d in docs]}

    # Global: return only files without course association
    from Coder.storage.course_manager import CourseManager
    docs = await CourseManager.list_global_files()
    return {"documents": [{
        "id": d["id"],
        "filename": d["filename"],
        "file_type": d["file_type"],
        "size": d["file_size"],
        "chunks": d["chunk_count"],
        "course_slug": "",
        "course_name": "全局知识库",
        "status": "indexed",
    } for d in docs]}


@router.delete("/documents/{file_id}")
async def delete_knowledge_document(file_id: str):
    """Delete a knowledge document: remove DB record, file on disk, and vector index."""
    from Coder.storage.course_manager import CourseManager

    info = await CourseManager.delete_file(file_id)
    if not info:
        raise HTTPException(status_code=404, detail="文档记录不存在")

    # Delete physical file
    docs_dir = os.path.join(os.path.dirname(__file__), "..", "..", "knowledge", "docs")
    file_path = os.path.join(docs_dir, info["filename"])
    try:
        os.remove(file_path)
    except FileNotFoundError:
        pass

    # Delete vector index files for this document's index path
    index_path = info.get("index_path", "")
    if index_path and os.path.isdir(index_path):
        import shutil
        shutil.rmtree(index_path)
        logger.info(f"已删除向量索引: {index_path}")

    # Delete associated knowledge points
    course_id = info.get("course_id")
    if course_id:
        await CourseManager.delete_knowledge_points_by_file(course_id, info["filename"])

    return {"status": "deleted", "filename": info["filename"]}


@router.get("/preview-docx")
async def preview_docx(path: str = ""):
    """预览 Word 文档内容。path 为文件名或相对于 workspace 的路径。"""
    import os as _os
    from fastapi import HTTPException

    workspace = _os.path.join(_os.path.dirname(__file__), "..", "..", "workspace")
    workspace = _os.path.normpath(_os.path.abspath(workspace))

    filename = _os.path.basename(path)
    if not filename.endswith(".docx"):
        filename += ".docx"
    filepath = _os.path.normpath(_os.path.join(workspace, filename))

    # 安全校验：必须在 workspace 内
    if not filepath.startswith(workspace):
        raise HTTPException(status_code=403, detail="路径不在 workspace 内")

    if not _os.path.exists(filepath):
        raise HTTPException(status_code=404, detail=f"文件不存在: {filename}")

    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(filepath)
        content = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else ""
            if style.startswith("Heading"):
                content.append({"type": "heading", "text": text, "level": 1 if "1" in style else 2})
            else:
                content.append({"type": "paragraph", "text": text})

        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append(rows)

        return {"filename": filename, "content": content, "tables": tables}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"读取文档失败: {e}")


def _resolve_docx_path(path: str) -> str:
    """解析 docx 文件的安全路径，返回绝对路径。"""
    import os as _os

    workspace = _os.path.join(_os.path.dirname(__file__), "..", "..", "workspace")
    workspace = _os.path.normpath(_os.path.abspath(workspace))

    filename = _os.path.basename(path)
    if not filename.endswith(".docx"):
        filename += ".docx"
    filepath = _os.path.normpath(_os.path.join(workspace, filename))

    if not filepath.startswith(workspace):
        from fastapi import HTTPException
        raise HTTPException(status_code=403, detail="路径不在 workspace 内")
    return filepath


@router.get("/docx-raw")
async def docx_raw(path: str = ""):
    """返回原始 .docx 文件二进制数据，供前端 mammoth.js 渲染。"""
    import os as _os
    from fastapi import HTTPException
    from fastapi.responses import FileResponse

    filepath = _resolve_docx_path(path)
    if not _os.path.exists(filepath):
        raise HTTPException(status_code=404, detail=f"文件不存在: {_os.path.basename(filepath)}")
    return FileResponse(filepath, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


