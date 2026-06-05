"""Word 文档工具 —— 生成和读取 .docx 文件。"""
import os
import logging
from langchain_core.tools import tool

logger = logging.getLogger(__name__)

_workspace = os.path.join(os.path.dirname(__file__), "..", "workspace")
os.makedirs(_workspace, exist_ok=True)


def _safe_path(filename: str) -> str:
    """防路径穿越，只允许在 workspace 内读写。"""
    filename = os.path.basename(filename)
    if not filename.endswith(".docx"):
        filename += ".docx"
    return os.path.normpath(os.path.join(_workspace, filename))


@tool
def create_docx(filename: str, title: str, content: str) -> str:
    """生成 Word 文档。

    filename: 文件名(不含路径，如 report.docx)
    title: 文档大标题
    content: 正文内容。用 ## 开头表示二级标题，用 --- 分隔段落，用 | 分隔表格列，用换行分隔表格行
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # 设置默认字体为微软雅黑，颜色为黑色
    _font_name = "微软雅黑"
    style = doc.styles["Normal"]
    style.font.name = _font_name
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), _font_name)
    # 标题样式也用微软雅黑 + 黑色
    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = _font_name
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), _font_name)

    # 标题
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 解析内容
    sections = content.split("---")
    for section in sections:
        section = section.strip()
        if not section:
            continue

        lines = section.split("\n")
        # 检测表格：第一行含 | 且至少两列
        if lines and "|" in lines[0]:
            cols = len(lines[0].split("|"))
            if cols >= 2:
                table = doc.add_table(rows=len(lines), cols=cols, style="Light Grid Accent 1")
                for i, line in enumerate(lines):
                    cells = [c.strip() for c in line.split("|")]
                    for j, cell_text in enumerate(cells):
                        if j < cols:
                            table.cell(i, j).text = cell_text
                continue

        # 二级标题
        if section.startswith("##"):
            doc.add_heading(section[2:].strip(), level=2)
            continue

        # 普通段落
        doc.add_paragraph(section)

    filepath = _safe_path(filename)
    doc.save(filepath)
    return f"Word 文档已生成: {filepath}"


@tool
def read_docx(filename: str) -> str:
    """读取 Word 文档内容。filename: 文件名(不含路径)。"""
    from docx import Document

    filepath = _safe_path(filename)
    if not os.path.exists(filepath):
        return f"文件不存在: {filepath}"

    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            level = para.style.name.split()[-1]
            prefix = "#" * int(level) if level.isdigit() else "##"
            parts.append(f"{prefix} {text}")
        else:
            parts.append(text)

    # 也读取表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n\n".join(parts)
